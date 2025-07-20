import os
import shutil
import uuid
import logging
import tempfile
import subprocess
import sys
from datetime import datetime
from pdf2docx import Converter
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import numpy as np
from docx import Document
import requests
import hashlib
import io

logger = logging.getLogger(__name__)

class PDFToWordConverter:
    def __init__(self):
        self.temp_dir = os.path.join(os.path.dirname(__file__), "temp")
        os.makedirs(self.temp_dir, exist_ok=True)
        self.ocr_enabled = True
        self.dpi = 300
        self.enhance_images = True
        self.remove_noise = True
        self.required_tools = {
            'tesseract': {
                'check': self.check_tesseract_installed(),
                'install': self.install_tesseract
            },
            'pdf2docx': True,
            'pymupdf': True
        }
        
    def check_tesseract_installed(self):
        """Check if Tesseract OCR is installed and accessible"""
        try:
            pytesseract.get_tesseract_version()
            return True
        except (pytesseract.TesseractNotFoundError, EnvironmentError):
            return False
    
    def check_requirements(self):
        """Check if all required tools are available"""
        missing = [name for name, data in self.required_tools.items() 
                 if not data.get('check', True)]
        return missing
    
    def install_tesseract(self):
        """Install Tesseract OCR based on the operating system"""
        try:
            if sys.platform == 'win32':
                return self.install_tesseract_windows()
            elif sys.platform == 'linux':
                return self.install_tesseract_linux()
            elif sys.platform == 'darwin':
                return self.install_tesseract_mac()
            return False
        except Exception as e:
            logger.error(f"Tesseract installation failed: {str(e)}")
            return False
    
    def install_tesseract_windows(self):
        """Install Tesseract OCR on Windows"""
        try:
            url = "https://digi.bib.uni-mannheim.de/tesseract/tesseract-ocr-w64-setup-5.3.1.20230401.exe"
            logger.info("Downloading Tesseract OCR for Windows...")
            installer_path = os.path.join(self.temp_dir, "tesseract_installer.exe")
            
            # Download the installer
            response = requests.get(url, stream=True)
            with open(installer_path, 'wb') as f:
                shutil.copyfileobj(response.raw, f)
            
            logger.info(f"Download complete. Please install Tesseract from: {installer_path}")
            logger.info("After installation, add Tesseract to your PATH environment variable")
            return False  # Return False because manual installation is needed
            
        except Exception as e:
            logger.error(f"Windows Tesseract installation failed: {str(e)}")
            return False
    
    def install_tesseract_linux(self):
        """Install Tesseract OCR on Linux"""
        try:
            # Try to install using apt-get
            result = subprocess.run(
                ['sudo', 'apt-get', 'install', '-y', 'tesseract-ocr'],
                capture_output=True,
                text=True
            )
            
            if result.returncode == 0:
                logger.info("Tesseract installed successfully on Linux")
                return True
            
            logger.error(f"Linux installation failed: {result.stderr}")
            return False
            
        except Exception as e:
            logger.error(f"Linux Tesseract installation failed: {str(e)}")
            return False
    
    def install_tesseract_mac(self):
        """Install Tesseract OCR on MacOS"""
        try:
            # Try to install using Homebrew
            result = subprocess.run(
                ['brew', 'install', 'tesseract'],
                capture_output=True,
                text=True
            )
            
            if result.returncode == 0:
                logger.info("Tesseract installed successfully on MacOS")
                return True
            
            logger.error(f"MacOS installation failed: {result.stderr}")
            return False
            
        except Exception as e:
            logger.error(f"MacOS Tesseract installation failed: {str(e)}")
            return False
    
    def generate_document_id(self, filename):
        """Generate unique document ID with timestamp and hash"""
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        unique_id = uuid.uuid4().hex[:8]
        file_hash = hashlib.md5(filename.encode()).hexdigest()[:6]
        return f"{timestamp}_{unique_id}_{file_hash}"
    
    def preprocess_image(self, image):
        """Enhance image quality for better OCR results"""
        try:
            # Convert to numpy array for processing
            img_array = np.array(image)
            
            # Contrast stretching
            p2, p98 = np.percentile(img_array, (2, 98))
            img_array = np.clip((img_array - p2) * 255.0 / (p98 - p2), 0, 255)
            
            if self.remove_noise:
                # Simple noise reduction using median blur
                img_array = cv2.medianBlur(img_array.astype(np.uint8), 3)
                
            # Convert back to PIL Image
            return Image.fromarray(img_array.astype(np.uint8)).convert('L')
            
        except Exception as e:
            logger.error(f"Image preprocessing failed: {str(e)}")
            return image.convert('L')  # Fallback to simple grayscale
    
    def extract_text_with_ocr(self, image, language='eng'):
        """Perform OCR on image with configurable language"""
        try:
            custom_config = r'--oem 3 --psm 6'
            if language != 'eng':
                custom_config += f' -l {language}'
                
            return pytesseract.image_to_string(
                image,
                config=custom_config
            )
        except Exception as e:
            logger.error(f"OCR failed: {str(e)}")
            return ""
    
    def add_text_to_docx(self, document, text, style=None):
        """Add formatted text to Word document"""
        para = document.add_paragraph()
        run = para.add_run(text)
        
        # Default styling
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        
        # Apply custom style if provided
        if style:
            if style.get('bold'):
                run.bold = True
            if style.get('italic'):
                run.italic = True
            if style.get('size'):
                run.font.size = Pt(style['size'])
    
    def convert_pdf_to_word(self, pdf_path, output_path=None, language='eng'):
        """
        Convert PDF to Word document with optional OCR
        Args:
            pdf_path: Path to input PDF file
            output_path: Optional output path for Word document
            language: Language for OCR (default: 'eng')
        Returns:
            tuple: (success: bool, result: str) where result is output path or error message
        """
        try:
            # Generate output path if not provided
            if not output_path:
                filename = os.path.basename(pdf_path)
                doc_id = self.generate_document_id(filename)
                output_path = os.path.join(
                    self.temp_dir,
                    f"{doc_id}.docx"
                )
            
            # First attempt: Use pdf2docx for direct conversion
            if not self.ocr_enabled:
                cv = Converter(pdf_path)
                cv.convert(output_path)
                cv.close()
                
                # Verify conversion quality
                if os.path.exists(output_path) and os.path.getsize(output_path) > 1024:
                    return True, output_path
            
            # Fallback to advanced processing
            pdf_document = fitz.open(pdf_path)
            docx_document = Document()
            
            # Set document properties
            docx_document.core_properties.title = f"Converted from {os.path.basename(pdf_path)}"
            docx_document.core_properties.author = "PDFMagic Converter"
            
            # Process each page
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Try to extract text directly first
                text = page.get_text()
                
                if text.strip():
                    # Direct text extraction worked
                    self.add_text_to_docx(docx_document, text)
                elif self.ocr_enabled:
                    # Use OCR for scanned pages
                    pix = page.get_pixmap(dpi=self.dpi)
                    img = Image.open(io.BytesIO(pix.tobytes()))
                    img = self.preprocess_image(img)
                    text = self.extract_text_with_ocr(img, language)
                    
                    if text.strip():
                        self.add_text_to_docx(docx_document, text)
            
            # Save the document
            docx_document.save(output_path)
            pdf_document.close()
            
            return True, output_path
            
        except Exception as e:
            logger.error(f"Conversion failed: {str(e)}", exc_info=True)
            return False, str(e)
    
    def convert_pdf_bytes_to_word(self, pdf_bytes, filename, language='eng'):
        """Convert PDF bytes to Word document in memory"""
        try:
            # Create temp PDF file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_temp:
                pdf_temp.write(pdf_bytes)
                pdf_path = pdf_temp.name
            
            # Convert to DOCX
            success, docx_path = self.convert_pdf_to_word(pdf_path, language=language)
            
            if success:
                with open(docx_path, "rb") as docx_file:
                    docx_bytes = docx_file.read()
                return True, docx_bytes
            
            return False, b""
            
        finally:
            # Cleanup temp files
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            if 'docx_path' in locals() and os.path.exists(docx_path):
                os.remove(docx_path)
    
    def cleanup_temp_files(self, max_age_hours=24):
        """Clean up old temporary files"""
        try:
            now = datetime.now()
            for filename in os.listdir(self.temp_dir):
                filepath = os.path.join(self.temp_dir, filename)
                file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                if (now - file_time).total_seconds() > max_age_hours * 3600:
                    try:
                        os.remove(filepath)
                        logger.info(f"Cleaned up old file: {filename}")
                    except Exception as e:
                        logger.error(f"Failed to clean up {filename}: {str(e)}")
        except Exception as e:
            logger.error(f"Temp file cleanup failed: {str(e)}")

# Helper function to check if file is a valid PDF
def is_pdf(filepath):
    try:
        with open(filepath, 'rb') as f:
            header = f.read(4)
            return header == b'%PDF'
    except:
        return False